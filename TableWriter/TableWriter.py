"""
Module that defines TableWriter class that writes mapping to docx table.
"""

import docx
from docx.text.run import Font, Run
from docx.dml.color import ColorFormat
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def shade(cell, shading):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), shading)
    table_cell_properties.append(shade_obj)


def add_text(cell, shading, text, italic, bold, style=None, alignment=None, color=None):
    paragraph = cell.paragraphs[0]
    paragraph.style = style if style is not None else paragraph.style
    paragraph.alignment = alignment if alignment is not None else paragraph.alignment
    runner = paragraph.add_run(text)
    runner.italic = italic if italic is not None else runner.italic
    runner.bold = bold if bold is not None else runner.bold
    if color is not None:
        font = runner.font
        font.color.rgb = RGBColor.from_string(color)

    if shading is not None:
        shade(cell, shading)


def change_borders(cell):
    color = 'bfbfbf'  # 'd9d9d9'
    val = 'single'
    cl = cell._tc
    tcPr = cl.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:color"), color)
    top.set(qn('w:val'), val)

    left = OxmlElement("w:left")
    left.set(qn("w:color"), color)
    left.set(qn('w:val'), val)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:color"), color)
    bottom.set(qn('w:val'), val)

    right = OxmlElement("w:right")
    right.set(qn("w:color"), color)
    right.set(qn('w:val'), val)

    tcBorders.append(top)
    tcBorders.append(left)
    tcBorders.append(bottom)
    tcBorders.append(right)
    tcPr.append(tcBorders)


class TableWriter:
    def __init__(self, title, filename):
        self.t = None
        self.title = title
        self.filename = filename
        # constants
        self.shadings = ['ffffff', 'ededed']
        self.bold = [True, False]
        self.doc = docx.Document('template.docx')
        self.init_table()

    def init_table(self, table_style='Table Grid'):
        # styling
        styles = self.doc.styles

        heading1 = styles['Heading 2']
        heading1.base_style = styles['Heading 2']
        font = heading1.font
        font.name = 'Calibri Light (Headings)'
        font.size = Pt(13)

        heading2 = styles.add_style('XPATH', WD_STYLE_TYPE.PARAGRAPH)
        heading2.base_style = styles['Normal']
        font = heading2.font
        font.name = 'Tahoma'
        font.size = Pt(9)

        normal = self.doc.styles['Normal']
        font = normal.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # add table
        self.t = self.doc.add_table(rows=0, cols=2, style='GridTable4-Accent31')

        # add heading
        row = self.t.add_row()
        a, b = row.cells
        a.merge(b)
        #add_text(a, 'a5a5a5', self.title, False, True, heading1, color='ffffff')
        add_text(a, None, self.title, False, True, heading1, color='ffffff')
        #change_borders(a)

        # add columns title
        row = self.t.add_row()
        #add_text(row.cells[0], self.shadings[1], 'XPATH/Element', True, True, heading2, 1)
        add_text(row.cells[0], None, 'XPATH/Element', True, True, heading2, 1)
        #change_borders(row.cells[0])

        #add_text(row.cells[1], self.shadings[1], 'Source Element (Expression)', True, False, heading2, 1)
        add_text(row.cells[1], None, 'Source Element (Expression)', True, False, heading2, 1)
        #change_borders(row.cells[1])

    def write_row(self, r):
        row = self.t.add_row()
        for j, c in enumerate(r):
            #add_text(row.cells[j], self.shadings[j | i % 2], c, False, self.bold[j])
            add_text(row.cells[j], None, c, None, None)
            #change_borders(row.cells[j])
        if len(r) == 1:
            a, b = row.cells
            a.merge(b)
            shade(a, 'a5a5a5')

    def save(self):
        self.doc.save(self.filename)